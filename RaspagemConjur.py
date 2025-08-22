import tkinter as tk
from tkinter import messagebox, filedialog
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import time
from docx import Document
from datetime import datetime
import pandas as pd


def buscar_noticias_conjur(tema, num_paginas=2):
    base_url = "https://www.conjur.com.br/pesquisa/"
    noticias = []

    options = webdriver.ChromeOptions()
    options.add_argument("--headless")
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)

    for pagina in range(1, num_paginas + 1):
        url = f"{base_url}?q={tema}&pagina={pagina}"

        try:
            driver.get(url)
            time.sleep(2)

            h2_elements = driver.find_elements(By.CSS_SELECTOR, 'h2 a')

            if not h2_elements:
                break

            for h2_element in h2_elements:
                titulo = h2_element.text
                link = h2_element.get_attribute('href')

                try:
                    data_element = h2_element.find_element(By.XPATH, '..//following-sibling::span')
                    data = data_element.text.strip() if data_element else "Data não encontrada"
                except Exception:
                    data = "Data não encontrada"

                noticias.append({
                    'titulo': titulo,
                    'data': data,
                    'link': link
                })

            time.sleep(1)

        except Exception as e:
            print(f"[ERRO] Falha ao acessar a página {pagina}: {e}")
            break

    driver.quit()
    return noticias


def salvar_noticias_em_doc(noticias, tema, caminho):
    doc = Document()
    doc.add_heading(f'Notícias ConJur - Tema: {tema}', 0)
    doc.add_paragraph(f'Data de Coleta: {datetime.now().strftime("%Y-%m-%d")}\n')

    for noticia in noticias:
        doc.add_paragraph(f"Título: {noticia['titulo']}")
        doc.add_paragraph(f"Data: {noticia['data']}")
        doc.add_paragraph(f"Link: {noticia['link']}")
        doc.add_paragraph("=" * 50)

    doc.save(caminho)


def salvar_noticias_em_xls(noticias, caminho):
    df = pd.DataFrame(noticias)
    df.to_excel(caminho, index=False)


def criar_interface():
    def iniciar_busca():
        tema_pesquisa = entry_tema.get()
        try:
            num_paginas = int(entry_paginas.get())
        except ValueError:
            messagebox.showerror("Erro", "O número de páginas deve ser um número inteiro.")
            return

        noticias = buscar_noticias_conjur(tema_pesquisa, num_paginas)

        if noticias:
            extensao = ".docx" if var_formato.get() == "docx" else ".xls"
            caminho_arquivo = filedialog.asksaveasfilename(
                defaultextension=extensao,
                filetypes=[("Documento Word", "*.docx"), ("Planilha Excel", "*.xls")],
                title="Salvar arquivo como"
            )

            if not caminho_arquivo:
                return

            if extensao == ".docx":
                salvar_noticias_em_doc(noticias, tema_pesquisa, caminho_arquivo)
            else:
                salvar_noticias_em_xls(noticias, caminho_arquivo)

            messagebox.showinfo("Sucesso", f"As notícias foram salvas em:\n{caminho_arquivo}")
        else:
            messagebox.showinfo("Resultado", "Nenhuma notícia encontrada.")

    root = tk.Tk()
    root.title("Busca de Notícias ConJur")

    tk.Label(root, text="Digite o tema da notícia:").pack(padx=20, pady=5)
    entry_tema = tk.Entry(root, width=50)
    entry_tema.pack(padx=20, pady=5)

    tk.Label(root, text="Número de páginas a ser pesquisado:").pack(padx=20, pady=5)
    entry_paginas = tk.Entry(root, width=10)
    entry_paginas.pack(padx=20, pady=5)

    # Opção de formato do arquivo
    frame_format = tk.Frame(root)
    frame_format.pack(pady=10)
    var_formato = tk.StringVar(value="docx")
    tk.Label(frame_format, text="Formato de arquivo: ").pack(side="left")
    tk.Radiobutton(frame_format, text="DOCX", variable=var_formato, value="docx").pack(side="left")
    tk.Radiobutton(frame_format, text="XLS", variable=var_formato, value="xls").pack(side="left")

    tk.Button(root, text="Buscar Notícias", command=iniciar_busca).pack(pady=20)

    root.mainloop()


criar_interface()
